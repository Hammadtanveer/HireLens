"""
engine/parser.py
------------------
Extracts plain text from uploaded resume files (.pdf, .docx, .txt) without
ever writing them to disk — everything is read from the in-memory upload
stream.
"""

import io
from pypdf import PdfReader
import docx


class UnsupportedFileType(Exception):
    pass


def extract_text(file_storage) -> str:
    """
    Takes a Werkzeug FileStorage object (from request.files) and returns
    its plain-text content, dispatching on file extension.
    """
    filename = (file_storage.filename or "").lower()
    raw = file_storage.read()
    file_storage.seek(0)  # reset stream in case it's needed again

    if filename.endswith(".pdf"):
        return _extract_pdf(raw)
    elif filename.endswith(".docx"):
        return _extract_docx(raw)
    elif filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"Unsupported file type: '{file_storage.filename}'. "
            "Please upload .pdf, .docx, or .txt files."
        )


def _extract_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(raw_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text out of any tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)
